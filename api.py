from bardapi import Bard
from flask import Flask, request
import os
import time
import docx
import PyPDF2
import pandas as pd
import google.generativeai as palm
from dotenv import load_dotenv
load_dotenv()
import requests
import io
import json
import cloudinary
from flask import jsonify
palm.configure(api_key=os.getenv("BARD_API_KEY"))
from langchain.text_splitter import CharacterTextSplitter
import base64
import cloudinary.uploader
stability_auth=os.getenv("STABILITY_AUTH")
from flask import send_file
app = Flask(__name__)
load_dotenv()
@app.route("/")
def hello_world():
    return "<p>Hello, World!</p>"
def get_pdf_text(file):
    text = ""
    pdfReader = PyPDF2.PdfFileReader(file)
    totalPages1 = pdfReader.numPages
    for i in range(totalPages1) :
        text = text+ pdfReader.pages[i].extract_text()

    return text

def get_text_from_csv(file):
    text = ""

    try:
        df = pd.read_csv(file, delim_whitespace=True)
        text = df.to_string(index=False)
    except Exception as e:
        print(f"Error reading CSV file: {e}")

    return text
def get_text_from_excel(file):
    text = ""

    try:
        df = pd.read_excel(file)
        text = df.to_string(index=False)
    except pd.errors.EmptyDataError:
        print("Excel file is empty.")
    except Exception as e:
        print(f"Error reading Excel file: {e}")

    return text
def get_text_from_docx(file):
    text = ""
    try:
        document = docx.Document(file)

        # abcc=document.paragraphs[1].text
        # print(abcc)
        for paragraph in document.paragraphs:
            text += paragraph.text + '\n'
    except Exception as e:
        print(f"Error reading Word file: {e}")

    return text
def get_file_text(file):
    text = ""
    if file.filename.endswith('.csv'):
        text = get_text_from_csv(file)
    elif file.filename.endswith('.xlsx'):
        text = get_text_from_excel(file)  # Assuming you have a function to handle Excel files
    elif file.filename.endswith('.docx'):
        text = get_text_from_docx(file)
    elif file.filename.endswith('.pdf'):
        text = get_pdf_text(file)
    else:
        # Handle other file types if needed
        pass

    return text
def split_text_chunks_and_summary_generator(text):
    text_splitter=CharacterTextSplitter(separator='\n',
                                        chunk_size=2500,
                                        chunk_overlap=20)
    text_chunks=text_splitter.split_text(text)
    print(len(text_chunks))
    return text_chunks
text_model="models/text-bison-001"
chat_model="models/chat-bison-001"
@app.route('/upload', methods=['POST'])
def upload_file():

  user_question = request.form['question']
  file = request.files['file']

  if file.filename == '':
        return 'No selected file'
  raw_text = get_file_text(file)
  text_chunks = split_text_chunks_and_summary_generator(raw_text)
  response_chunks = []
  for chunk in text_chunks:
        response_chunk = palm.generate_text(prompt="write a summery of "+chunk+"in"+str(3000/(len(text_chunks)+1))+"Word", temperature=0.1, model=text_model)
        response_chunks.append(response_chunk)
  combined_response = ""
  for response_chunk in response_chunks:
        combined_response += str(response_chunk.result) + "\n"


  return combined_response
def ocr_space_file(filename, overlay=False, api_key=os.getenv("OCR_KEY"), language='eng'):

    payload = {'isOverlayRequired': overlay,
               'apikey': api_key,
               'language': language,
               }


    with open(filename, 'rb') as f:
        r = requests.post('https://api.ocr.space/parse/image',
                          files={filename: f},

                          data=payload,
                          )
    return r.content.decode()
@app.route('/ocr', methods=['POST'])
def ocr():
    image = request.files['file']
    image.save("images/filename.png")
    a=ocr_space_file("images/filename.png")
    b=json.loads(a)
    return b["ParsedResults"][0]["ParsedText"]


@app.route('/url', methods=['POST'])
def fromurl():
    url = request.form['url']
    response = palm.generate_text(prompt="give me brief about content in this url"+url, temperature=0.1, model=text_model)
    return response.result



@app.route('/qna', methods=['POST'])
def qna():
    next_question = request.form['question']
    text=request.form['text']


    thisr=palm.chat(messages="ans the question on these context\n"+text, temperature=0.1, model=chat_model)
    thisr = thisr.reply(next_question)

    responce_data={
    "question": next_question,
    "ans":thisr.last,
    "text":text
    }
    return jsonify(responce_data)

def upload_file():
  app.logger.info('in upload route')

  cloudinary.config(cloud_name = os.getenv("CLOUDINARY_CLOUD_NAME"), api_key=os.getenv("CLOUDNARY_API_KEY"),
    api_secret=os.getenv("CLOUDINARY_SECRET"))
  upload_result = None
  if request.method == 'POST':
    file_to_upload = request.files['file']
    app.logger.info('%s file_to_upload', file_to_upload)
    if file_to_upload:
      upload_result = cloudinary.uploader.upload(file_to_upload)
      app.logger.info(upload_result)
      return jsonify(upload_result)

@app.route('/generateimage', methods=['POST'])
def generateimage():
    prompt = request.form['prompt']
    url = "https://api.stability.ai/v1/generation/stable-diffusion-xl-1024-v1-0/text-to-image"
    body = {
  "steps": 10,
  "width": 1024,
  "height": 1024,
  "seed": 0,
  "cfg_scale": 5,
  "samples": 1,
  "text_prompts": [
    {
      "text": prompt,
      "weight": 1
    },
    {
      "text": "blurry, bad",
      "weight": -1
    }
  ],
}

    headers = {
  "Accept": "application/json",
  "Content-Type": "application/json",
  "Authorization": "Bearer ${stability_auth}",
}
    response = requests.post(
    url,
    headers=headers,
    json=body,
    )
    data = response.json()

    for i, image in enumerate(data["artifacts"]):
        with open('abc.png', "wb") as f:
            f.write(base64.b64decode(image["base64"]))
    # cloudinary.config(cloud_name = "dm77rbalw", api_key="538483726846567",
    # api_secret="pN4v-PM_tmbzdirmEkcqYU19hCI",api_proxy= 'http://proxy.server:3128')
    # res =cloudinary.uploader.upload("abc.png")

    return send_file('abc.png', mimetype='image/gif')
@app.route('/imagerecognition', methods=['POST'])
def ImageRecognition():
    imageurl=request.form['imageurl']
    api_key = os.getenv("imagga_api_key")
    api_secret = os.getenv("imagga_api_secret")

    response = requests.get(
        'https://api.imagga.com/v2/tags?image_url=%s' % imageurl,
        auth=(api_key, api_secret))
    q=response.json()

    z=q['result']['tags'][0]['tag']['en']

    # for x in range(5):
    #     rec.append(q['result']['tags'][x]['tag']['en'])
    return z
@app.route('/chat', methods=['POST'])
def chat():
    # next_question = request.form['question']
    history=request.form['history']
    prompt=request.form['prompt']
    bard = Bard(token=os.getenv("BARD_TOKEN"))


    thisr=bard.get_answer(f"Conversation History: {history}\n {prompt}")
    # thisr = thisr.reply(next_question)

    # responce_data={
    # "question": next_question,
    # "ans":thisr.last,
    # "text":text
    # }
    return thisr["content"]

if __name__ == "__main__":
    app.run()